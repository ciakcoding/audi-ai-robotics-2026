"""Fill sections 5-10 tables + section 11 + guided questions + checklist."""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

doc = Document("document/Sim2Real_ADC2026_v2.docx")

# ── Table 5: Problems & Setbacks (cols: Problem | Root Cause | Fix | Result | Status) ──
problems = [
    ["Ball mass masked real findings",
     "Initial tests included ball mass randomization (0.06-0.12 kg) in the perturbation set",
     "Removed ball_mass and ball_size from randomization; fixed ball at 50g/4cm",
     "Clean results: remaining 7 perturbations show 97% success vs 100% clean",
     "Resolved"],
    ["Per-parameter isolation was time-consuming",
     "Combined test showed 0.039m error but did not explain which parameter caused it",
     "Ran 7 separate single-parameter tests (V2-V8) at 30 episodes each",
     "Identified target_noise as the only parameter with solo impact (4.2cm vs ~2cm others)",
     "Resolved"],
    ["Small sample size limits confidence",
     "30-episode batches cannot reliably distinguish 3% success differences",
     "Use 100-episode batches for combined tests; accept 30-episode for per-parameter screening",
     "Combined test (100ep): 97% success confirmed. Per-parameter (30ep): trends clear but not definitive",
     "Acknowledged"],
    ["Target position noise emerged as top risk",
     "Template Gap table originally treated all 7 gaps equally",
     "Re-prioritized after per-parameter testing showed target_noise as the only high-impact gap",
     "Gap Analysis table updated: target position = High, all others = Low/Medium",
     "Resolved"],
    ["Teammate repository structure mismatch",
     "Teammate repo uses residual PPO (ppo_throw_env.py) while teacher template assumes pure PPO",
     "Adapted robustness env to inherit from teammate's PPOThrowEnv rather than G1FixedBodyThrowEnv",
     "Robustness tests work correctly with teammate's best_model.zip",
     "Resolved"],
]
t5 = doc.tables[4]
for ri, row_data in enumerate(problems):
    row = t5.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].text = val

# ── Table 6: Engineering Decisions (cols: Decision | Reason | Alternative | Impact) ──
decisions = [
    ["Fixed ball parameters (50g, 4cm)",
     "Ball mass variation was the dominant noise source (43% success drop alone). The real test setup should use a precisely measured ball.",
     "Train with wider ball mass range to produce mass-robust policy",
     "Enabled clean isolation of robot-related Sim2Real gaps. Policy is 97% robust with ball fixed."],
    ["Per-parameter isolation testing",
     "Combined testing only shows total degradation; isolation reveals which gaps matter and which do not.",
     "Only run combined test and report aggregate degradation without understanding causes",
     "Identified target_noise as the sole high-impact gap. Saved effort on 6 irrelevant perturbation types."],
    ["Residual PPO over pure PPO",
     "Team trained PPO as correction on scripted baseline. More stable, faster convergence, maintains 100% baseline success.",
     "Pure PPO from scratch (teacher's train_g1_throw.py) — higher ceiling but more training risk",
     "Policy is stable (100% baseline success) but improvement over baseline is limited to ~41% error reduction."],
    ["10cm success radius",
     "Matches target geom size in scene_throw.xml. Conservative enough to credit clear successes, tight enough to reject failures.",
     "Smaller radius (5cm) for higher precision requirement; larger (15cm) for easier success",
     "All tests use consistent 10cm criterion. Baseline error (1.15cm) and PPO error (0.67cm) both well inside."],
]
t6 = doc.tables[5]
for ri, row_data in enumerate(decisions):
    row = t6.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].text = val

# ── Table 7: Limitations (cols: Limitation | Impact | Possible Solution) ──
limitations = [
    ["Fixed body only (right arm, 7 DoF)",
     "Cannot demonstrate whole-body stability. Real G1 must support itself or be externally held.",
     "Unlock legs and torso joints in training; add balance reward terms."],
    ["Scripted release time (t=0.65s)",
     "Policy cannot adapt release timing. If arm trajectory deviates, ball launches from wrong pose.",
     "Add learned release action to PPO action space."],
    ["No real hardware validation",
     "All results are simulation-only. True Sim2Real gap is unmeasured.",
     "Deploy on physical Unitree G1 when hardware access is available."],
    ["Single target position (0.55, 0, 0)",
     "Policy is not tested on different drop locations. Generalization is unknown.",
     "Train with randomized target positions; evaluate on held-out locations."],
    ["30-episode per-parameter batches",
     "3% success differences are at edge of statistical significance with n=30.",
     "Use 100-episode batches for all confirmatory tests."],
    ["Worst-case outlier (~20cm in V1)",
     "Despite 97% mean success, single large-error episodes indicate fragile edge cases.",
     "Log and replay worst-case episodes to identify failure conditions."],
]
t7 = doc.tables[6]
for ri, row_data in enumerate(limitations):
    row = t7.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].text = val

# ── Table 8: Future Work (cols: Item | Expected Benefit | Priority) ──
future = [
    ["Whole-body training",
     "Policy controls full 29-DoF G1; robot can use legs and torso for more natural ball-dropping motion",
     "Medium"],
    ["Learned ball release",
     "Policy adapts release timing to arm state; more robust to trajectory variations",
     "High"],
    ["Multi-target generalization",
     "Policy works for arbitrary drop positions; enables pick-and-place style tasks",
     "Medium"],
    ["Domain randomization during training",
     "Inherently robust policy; no separate post-training robustness testing needed",
     "High"],
    ["Wider randomization ranges",
     "Identify true failure boundaries for each parameter; safety margins for deployment",
     "Low"],
    ["Real G1 hardware deployment",
     "Close the Sim2Real loop; measure actual vs simulated landing error",
     "High"],
]
t8 = doc.tables[7]
for ri, row_data in enumerate(future):
    row = t8.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].text = val

# ── Table 9: Evidence Register (cols: Evidence Item | Filename | Description | Presentation?) ──
evidence = [
    ["Robustness combined (100ep)",
     "outputs/robustness_final.json",
     "Clean vs noisy: 100 episodes each, 7 perturbations active",
     "Yes"],
    ["Per-parameter isolation (7x30ep)",
     "outputs/per_param_results.json",
     "V2-V8 single-parameter tests identifying target_noise as dominant factor",
     "Yes"],
    ["Teammate training report",
     "teammate_repo/docs/TRAINING_REPORT.md",
     "3 completed 1M-step PPO runs, reward iteration, model selection rationale",
     "Yes"],
    ["Nominal comparison",
     "teammate_repo/outputs/plots/nominal/summary.json",
     "Baseline vs PPO best vs PPO final: 100-episode frozen-world protocol",
     "Yes"],
    ["Robustness comparison",
     "teammate_repo/outputs/plots/robustness/summary.json",
     "Same 100 seeds with +0.08rad joint noise for both policies",
     "Yes"],
    ["Scene definition",
     "assets/scene_throw.xml",
     "Ball 50g/4cm, target (0.55,0,0), success radius 0.10m",
     "No"],
    ["Robustness environment",
     "envs/g1_robustness_env.py",
     "7 domain randomization parameters, per-parameter control",
     "No"],
    ["Evaluation script",
     "scripts/evaluate_robustness.py",
     "Clean vs noisy comparison with configurable perturbation set",
     "No"],
]
t9 = doc.tables[8]
for ri, row_data in enumerate(evidence):
    row = t9.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].text = val

# ── Table 10: Presentation Assets (cols: Asset | Filename | Key Message) ──
assets = [
    ["Side-by-side viewer",
     "scripts/compare_side_by_side.py",
     "Dual MuJoCo windows: clean (left) vs 7 perturbations (right) — visual proof of robustness"],
    ["Per-parameter bar chart",
     "Generate from outputs/per_param_results.json",
     "Target_noise is the only high-impact gap; all others have negligible solo effect"],
    ["CDF landing error plot",
     "Generate from outputs/robustness_final.json",
     "Distribution shift from tight 0.02m cluster (clean) to 0.02-0.20m spread (noisy)"],
    ["Training reward curve",
     "teammate_repo/outputs/plots/training_evaluation_curve.png",
     "PPO convergence over 1M steps; best model selected at peak evaluation reward"],
    ["Baseline vs PPO bar chart",
     "teammate_repo/outputs/plots/nominal/baseline_vs_ppo.png",
     "PPO reduces mean error by 41% vs baseline (1.15cm -> 0.67cm)"],
    ["Robustness comparison chart",
     "teammate_repo/outputs/plots/robustness/baseline_vs_ppo.png",
     "PPO maintains advantage under joint noise: 0.75cm vs baseline 1.18cm"],
]
t10 = doc.tables[9]
for ri, row_data in enumerate(assets):
    row = t10.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].text = val

# ── Section 11: Lessons Learned (replace paragraph text) ──
body = doc.element.body
children = list(body)
for child in children:
    if child.tag != qn("w:p"):
        continue
    texts = child.findall(".//" + qn("w:t"))
    full = "".join(t.text or "" for t in texts)
    if "11. Lessons" in full:
        # Find next paragraph (which should be empty or the content placeholder)
        idx = list(body).index(child)
        # Check if there's already a content paragraph after
        next_p = None
        for j in range(idx + 1, min(idx + 3, len(children))):
            if children[j].tag == qn("w:p"):
                ptexts = children[j].findall(".//" + qn("w:t"))
                pfull = "".join(t.text or "" for t in ptexts)
                if not any(kw in pfull for kw in ["12.", "Sim2Real"]):
                    next_p = children[j]
                    break
        if next_p is not None:
            for r in next_p.findall(qn("w:r")):
                next_p.remove(r)
            content = (
                "Lesson 1 - Always isolate parameters: Combined perturbation testing tells you performance degrades, "
                "but only per-parameter isolation tells you what to fix. Our initial all-7 test showed 0.039m error; "
                "isolation revealed target_noise alone caused 0.042m while 6 other parameters had zero individual impact.\n\n"
                "Lesson 2 - Fix what matters, do not randomize everything: Including ball mass in the perturbation set "
                "masked all other findings and wasted the first round of testing. Define scope clearly (ball is fixed) "
                "before designing tests.\n\n"
                "Lesson 3 - The policy was more robust than expected: With ball fixed at 50g/4cm, 7 simultaneous "
                "perturbations caused only a 3% success drop (100% -> 97%). The residual PPO approach (baseline + "
                "learned corrections) produces inherently stable behavior.\n\n"
                "Lesson 4 - Target position is the real bottleneck: In a real lab, the biggest risk is not sensor noise "
                "or motor degradation, but simply not knowing exactly where the target is. A clearly marked, precisely "
                "measured target zone is the cheapest and most effective robustness improvement.\n\n"
                "Lesson 5 - Statistical significance matters: 30 episodes can detect large effects (>10% success change) "
                "but small differences (3%) need 100+ episodes. For the final combined test, we used 100-episode runs."
            )
            for line in content.split("\n"):
                r = etree.SubElement(next_p, qn("w:r"))
                t = etree.SubElement(r, qn("w:t"))
                t.text = line
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        break

# ── Guided questions ──
for child in children:
    if child.tag != qn("w:p"):
        continue
    texts = child.findall(".//" + qn("w:t"))
    full = "".join(t.text or "" for t in texts)
    if "Which Sim2Real gap was the most critical?" in full:
        for r in child.findall(qn("w:r")):
            child.remove(r)
        content = (
            "Which Sim2Real gap was the most critical?\n"
            "Target position uncertainty. When the target position was randomized by +/-3cm, "
            "mean landing error doubled from 2.0cm to 4.2cm and success rate dropped from 100% to 97%. "
            "All other gaps (sensor noise, friction, latency, contact, motor dynamics) had negligible "
            "individual impact (~2cm, 100% success). The single best investment for real-world deployment "
            "is a precisely measured and clearly marked target zone.\n\n"
            "Which robustness technique appears most promising?\n"
            "Per-parameter isolation testing. Rather than only measuring combined degradation, isolating "
            "each parameter revealed that 6 out of 7 perturbations are non-issues individually. This allows "
            "us to focus engineering effort on the one real problem (target position) instead of wasting time "
            "on sensor noise filtering or actuator modeling improvements that would have zero impact.\n\n"
            "What would you do differently?\n"
            "1. Define ball as fixed parameter from the start (50g/4cm measured).\n"
            "2. Run per-parameter tests before combined tests to understand individual contributions first.\n"
            "3. Use 100-episode batches for all confirmatory tests, not just the combined one.\n"
            "4. Add target position randomization to the training loop (domain randomization during PPO "
            "training) rather than only at evaluation time."
        )
        for line in content.split("\n"):
            r = etree.SubElement(child, qn("w:r"))
            t = etree.SubElement(r, qn("w:t"))
            t.text = line
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        break

# ── Checklist ──
checklist_map = {
    "Main transfer gaps identified": "[X] Main transfer gaps identified",
    "Domain randomization documented": "[X] Domain randomization documented",
    "Robustness experiments completed": "[X] Robustness experiments completed",
    "Limitations documented": "[X] Limitations documented",
    "Future work proposed": "[X] Future work proposed",
    "Evidence collected for presentation": "[X] Evidence collected for presentation",
}
for child in children:
    if child.tag != qn("w:p"):
        continue
    texts = child.findall(".//" + qn("w:t"))
    full = "".join(t.text or "" for t in texts)
    for key, val in checklist_map.items():
        if key in full:
            for r in child.findall(qn("w:r")):
                child.remove(r)
            r = etree.SubElement(child, qn("w:r"))
            t = etree.SubElement(r, qn("w:t"))
            t.text = val
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            break

doc.save("document/Sim2Real_ADC2026_v2.docx")
print("DONE: All sections filled with proper table structure")
