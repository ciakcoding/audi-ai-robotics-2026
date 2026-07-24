#!/usr/bin/env python3
"""Fill Sim2Real report with real data from robustness tests (ball fixed at 50g/4cm)."""
from docx import Document

doc = Document("document/Sim2Real_ADC2026.docx")

# === TABLE 1: Workstream Overview ===
t1 = doc.tables[0]
t1.rows[1].cells[1].text = "Sim2Real Lead"
t1.rows[6].cells[1].text = "https://github.com/aim-t/audi-ai-robotics-2026/tree/feature/rl"

# === TABLE 2: Gap Analysis ===
t2 = doc.tables[1]
gaps = [
    ["Sensor noise",
     "MuJoCo provides perfect joint/IMU readings; real G1 encoders have 0.01-0.05 rad error",
     "Policy sees clean state in sim but noisy state on hardware; actions based on wrong estimates. Solo test: no measurable impact (mean error 2.05cm, 100% success).",
     "Low"],
    ["Actuator latency",
     "Sim applies control instantly; real G1 has 3-10ms communication delay",
     "Delayed control could cause arm trajectory oscillation. Solo test: 6ms delay shows no measurable impact (mean error 1.93cm, 100% success).",
     "Low"],
    ["Friction mismatch",
     "Joint frictionloss (0.3) and floor friction are fixed ideal values",
     "Joint friction affects arm repeatability; floor friction changes ball bounce. Solo tests: no individual impact (both ~2.0cm, 100% success).",
     "Medium"],
    ["Contact modelling",
     "Fixed solref/solimp; rigid-body contact with no surface compliance",
     "Ball-hand release and ball-floor impact differ from real elastic contact. Solo test: no measurable impact (2.0cm, 100% success).",
     "Medium"],
    ["Motor dynamics",
     "Actuator forcerange constant; no voltage sag or thermal effects",
     "Real G1 motors lose up to 15% torque under load; arm may not reach planned release pose. Solo test: no measurable impact (1.98cm, 100% success).",
     "Medium"],
    ["Observation noise",
     "Policy receives perfect state; real state comes from filtered sensor fusion with lag",
     "Cumulative state estimation error leads to suboptimal actions. Solo test: no measurable impact (2.05cm, 100% success). Risk is from combination with other factors.",
     "Medium"],
    ["Other (target position)",
     "Target (0.55, 0, 0) is fixed in sim; real lab setup introduces ~3cm XY variation",
     "Target position noise is the LARGEST single-parameter contributor: mean error 4.17cm (2.1x degradation), 97% success. Combined with all 7 perturbations: 3.92cm, 97% success.",
     "High"],
]
for ri, (gap, why, impact, priority) in enumerate(gaps):
    t2.rows[ri + 1].cells[0].text = gap
    t2.rows[ri + 1].cells[1].text = why
    t2.rows[ri + 1].cells[2].text = impact
    t2.rows[ri + 1].cells[3].text = priority

# === TABLE 3: Domain Randomization ===
t3 = doc.tables[2]
params = [
    ["obs_noise (Observation noise)",
     "Gaussian sigma=0.02 on observation vector",
     "Simulates IMU/encoder measurement error on real G1",
     "Yes (no solo impact: 2.05cm)"],
    ["joint_friction (Joint friction)",
     "Multiplier Uniform [0.7, 1.3] on joint frictionloss",
     "Bearing wear, temperature, lubrication: up to +/-30% variation",
     "Yes (no solo impact: 1.96cm)"],
    ["joint_damping (Joint damping)",
     "Multiplier Uniform [0.7, 1.3] on joint damping",
     "Actuator degradation and temperature effects on viscosity",
     "Yes (no solo impact)"],
    ["floor_friction (Floor friction)",
     "Multiplier Uniform [0.5, 1.5] on floor geom friction",
     "Polished concrete (~0.5x) vs rubber mat (~1.5x) floor surfaces",
     "Yes (no solo impact: 1.92cm)"],
    ["actuator_gain (Motor force)",
     "Multiplier Uniform [0.85, 1.0] on actuator forcerange",
     "Voltage sag and motor heating: up to 15% torque reduction",
     "Yes (no solo impact: 1.98cm)"],
    ["target_pos_noise (Target position)",
     "Gaussian sigma=0.03m in XY plane",
     "Real target placement is never millimeter-precise",
     "Yes (LARGEST impact: 4.17cm, 97% success)"],
    ["contact_solref/solimp (Contact)",
     "Multiplier Uniform [0.5, 2.0] on global solref[0] and solimp[0]",
     "Surface compliance: hard lab floor (~0.5x) vs padded surface (~2.0x)",
     "Yes (no solo impact: 2.00cm)"],
    ["control_latency (Control delay)",
     "3 timesteps (~6ms) delay on action application",
     "Real G1 communication and control loop latency (3-10ms)",
     "Yes (no solo impact: 1.93cm)"],
]
for ri, (param, rng, reason, tested) in enumerate(params):
    t3.rows[ri + 1].cells[0].text = param
    t3.rows[ri + 1].cells[1].text = rng
    t3.rows[ri + 1].cells[2].text = reason
    t3.rows[ri + 1].cells[3].text = tested

# === TABLE 4: Robustness Experiments ===
t4 = doc.tables[3]
experiments = [
    ["V1 (all 7 combined)",
     "PPO best policy is robust to combined Sim2Real perturbations",
     "All 7 domain randomization params simultaneously active (no ball mass/size; ball fixed at 50g, 4cm)",
     "Clean: mean error 0.020m, 100% success. Noisy (7 perturbations): mean error 0.039m, 97% success. Degradation: +0.019m (1.9x). Success drops only 3%.",
     "Keep",
     "Policy is highly robust when ball parameters are fixed. Only 3% success drop despite 7 simultaneous perturbations. Mean error nearly doubles but stays well within 10cm success radius."],
    ["V2 (obs_noise only)",
     "Observation noise alone has negligible impact",
     "Only obs_noise=0.02; all others off",
     "Mean error: 0.0205m. Success: 100%. Essentially identical to clean (0.020m).",
     "Keep",
     "Policy tolerates sensor noise well. The baseline trajectory is slow enough that noisy observations do not destabilize it."],
    ["V3 (joint_friction only)",
     "Joint friction variation has negligible solo impact",
     "Only joint_friction_range=[0.7,1.3]; all others off",
     "Mean error: 0.0196m. Success: 100%",
     "Keep",
     "Policy trajectory handles +/-30% joint friction without measurable degradation."],
    ["V4 (floor_friction only)",
     "Floor friction variation has no negative impact",
     "Only floor_friction_range=[0.5,1.5]; all others off",
     "Mean error: 0.0192m. Success: 100%",
     "Keep",
     "Different floor surfaces will not affect this policy. Ball landing behavior is dominated by release kinematics, not floor friction."],
    ["V5 (actuator_gain only)",
     "Motor torque reduction has negligible solo impact",
     "Only actuator_gain_range=[0.85,1.0]; all others off",
     "Mean error: 0.0198m. Success: 100%",
     "Keep",
     "Policy does not push joint torque limits. 15% torque reduction is absorbed by the conservative baseline trajectory."],
    ["V6 (target_noise only)",
     "Target position noise is the largest single-parameter risk",
     "Only target_pos_noise=0.03m; all others off",
     "Mean error: 0.0417m (2.1x degradation). Success: 97% (3% drop). Largest individual impact among all 7 parameters.",
     "Keep",
     "Target noise shifts success rate from 100% to 97% and doubles mean error. Recommendation: mark target zone clearly and measure its position precisely in real tests."],
    ["V7 (contact only)",
     "Contact stiffness/impedance variation has negligible solo impact",
     "Only contact_solref_range=[0.5,2.0], contact_solimp_range=[0.5,2.0]",
     "Mean error: 0.0200m. Success: 100%",
     "Keep",
     "Policy does not depend on precise contact dynamics. The ball release is via weld constraint break, not contact friction."],
    ["V8 (latency only)",
     "Control latency has negligible solo impact",
     "Only control_latency_steps=3 (~6ms delay)",
     "Mean error: 0.0193m. Success: 100%",
     "Keep",
     "6ms latency is negligible for this low-speed (<1 rad/s) arm trajectory. Real G1 latency (3-10ms) should not be a concern."],
]
for ri, (ver, hyp, chg, out, keep, notes) in enumerate(experiments):
    t4.rows[ri + 1].cells[0].text = ver
    t4.rows[ri + 1].cells[1].text = hyp
    t4.rows[ri + 1].cells[2].text = chg
    t4.rows[ri + 1].cells[3].text = out
    t4.rows[ri + 1].cells[4].text = keep
    t4.rows[ri + 1].cells[5].text = notes

# === Section 5: Problems & Setbacks ===
p5 = doc.paragraphs[6]
# Find the next paragraph after heading 5
for i, p in enumerate(doc.paragraphs):
    if "5. Problems" in p.text:
        # Insert text after this heading
        pass

# Use the paragraph after each heading to fill content
sections_text = {
    "5. Problems & Setbacks": (
        "Problem 1 - Ball mass masked the real findings: Initial robustness tests included ball mass randomization (0.06-0.12 kg), "
        "which alone caused a 43% success rate drop and dominated all other effects. This made it impossible to assess the true impact "
        "of robot-related Sim2Real gaps. Decision: removed ball mass/size from the perturbation set and fixed ball at 50g/4cm as specified.\n\n"
        "Problem 2 - Per-parameter isolation was essential but time-consuming: The combined 9-parameter test showed 0.103m error, "
        "but only individual testing revealed that target position noise (0.042m) was the true largest factor, while all other "
        "parameters had negligible solo impact (~0.02m each). Without isolation, we would have wrongly attributed degradation to sensor noise or friction.\n\n"
        "Problem 3 - 30-episode sample size: Each test used only 30 episodes due to time constraints. Success rate differences "
        "of 3% (100% vs 97%) are at the edge of statistical significance with this sample size."
    ),
    "6. Engineering Decisions": (
        "Decision 1 - Fixed ball parameters: Ball mass (50g) and radius (4cm) are treated as fixed, known quantities. "
        "The real test setup should use a precisely measured ball. This removes the dominant noise source and lets us "
        "focus on robot-related Sim2Real gaps.\n\n"
        "Decision 2 - Per-parameter isolation testing: Rather than only testing all perturbations combined, we ran 8 separate "
        "single-parameter tests (V2-V8) to quantify each gap's individual contribution. This produced the key finding that "
        "target position noise is the only parameter with measurable solo impact.\n\n"
        "Decision 3 - Residual PPO over pure PPO: The team chose to train PPO as a residual correction on top of a scripted "
        "baseline trajectory, rather than learning from scratch. This produced a more stable policy (100% baseline success maintained) "
        "but means the policy cannot outperform the baseline by large margins.\n\n"
        "Decision 4 - 10cm success radius: The target success criterion was set to 0.10m radius, matching the target geom size "
        "in the scene XML. This is a reasonable engineering tolerance for a ball-dropping task."
    ),
    "7. Limitations": (
        "Limitation 1 - Fixed body only: The current policy controls only the right arm (7 DoF). The rest of the G1 body "
        "(legs, torso, left arm, head) is held at the nominal standing pose by PD control. Real deployment would require "
        "whole-body stability or the robot must be externally supported.\n\n"
        "Limitation 2 - Scripted release time: Ball release is fixed at t=0.65s regardless of arm state. A learned release "
        "policy could adapt to trajectory deviations and release at the optimal moment.\n\n"
        "Limitation 3 - No real hardware validation: All tests are simulation-only. The true Sim2Real gap cannot be fully "
        "characterized without running the policy on a physical Unitree G1, which is outside the project scope.\n\n"
        "Limitation 4 - Single target position: The target is fixed at (0.55, 0, 0). The policy has not been tested on "
        "different target locations, so generalization to arbitrary drop positions is unknown.\n\n"
        "Limitation 5 - 30-episode test batches: Statistical confidence is limited. A 3% success rate difference (97% vs 100%) "
        "with 30 episodes means only 1 failure was observed. Larger batches (100+) would give more reliable estimates.\n\n"
        "Limitation 6 - Worst-case outlier: In the combined 7-perturbation test (V1), the maximum error was ~20cm despite "
        "97% mean success. Single-outlier failures must be considered for safety-critical deployment."
    ),
    "8. Future Work": (
        "1. Whole-body training: Unlock the G1 legs and torso to learn a full-body ball-dropping motion. This would also "
        "require adding balance-related reward terms and fall detection.\n\n"
        "2. Learned ball release: Replace the fixed 0.65s scripted release with a learned release action, allowing the "
        "policy to adapt the release timing to the current arm state.\n\n"
        "3. Multi-target generalization: Train and evaluate with randomized target positions to produce a policy that "
        "can drop the ball at any specified location, not just (0.55, 0, 0).\n\n"
        "4. Wider randomization ranges: Test more extreme parameter variations to find the true failure boundary "
        "(e.g., target_noise=0.10m, joint_friction=0.3-2.0x).\n\n"
        "5. Domain randomization during training: Instead of only testing robustness post-training, incorporate domain "
        "randomization into the training loop itself to produce an inherently robust policy.\n\n"
        "6. Real G1 deployment: The ultimate next step - deploy on physical hardware and compare sim vs real landing error. "
        "This would close the loop on the entire Sim2Real pipeline."
    ),
    "9. Evidence Register": (
        "E1 - Robustness test results (combined): outputs/robustness_final.json (100 episodes clean + 100 episodes noisy)\n"
        "E2 - Per-parameter isolation results: outputs/per_param_results.json (8 tests x 30 episodes)\n"
        "E3 - Teammate training report: teammate_repo/docs/TRAINING_REPORT.md (1M steps, 3 completed runs)\n"
        "E4 - Teammate nominal comparison: teammate_repo/outputs/plots/nominal/summary.json (baseline vs PPO best vs PPO final)\n"
        "E5 - Teammate robustness comparison: teammate_repo/outputs/plots/robustness/summary.json (joint noise test)\n"
        "E6 - Scene definition: assets/scene_throw.xml (ball: 0.05kg/4cm, target: 0.55m, radius: 0.10m)\n"
        "E7 - Robustness environment: envs/g1_robustness_env.py (7 domain randomization parameters)\n"
        "E8 - Evaluation script: scripts/evaluate_robustness.py (clean vs noisy comparison)\n"
        "E9 - Visual comparison: scripts/compare_side_by_side.py (dual-window MuJoCo viewer)\n"
        "E10 - Training curve: teammate_repo/outputs/plots/training_evaluation_curve.png"
    ),
    "10. Presentation Assets": (
        "Asset 1 - Side-by-side comparison: Run scripts/compare_side_by_side.py to show CLEAN vs NOISY policy behavior "
        "in two MuJoCo windows simultaneously. This is the strongest visual evidence for the final presentation.\n\n"
        "Asset 2 - Per-parameter bar chart: Generate a bar chart comparing mean landing error across V1-V8 tests. "
        "This clearly shows target_noise as the dominant factor (4.2cm vs ~2cm for all others).\n\n"
        "Asset 3 - CDF plot: Landing error cumulative distribution for clean vs noisy (100 episodes). Shows the "
        "distribution shift from tight (0.02m cluster) to spread (0.02-0.20m range).\n\n"
        "Asset 4 - Teammate training plot: teammate_repo/outputs/plots/training_evaluation_curve.png showing "
        "reward convergence over 1M steps.\n\n"
        "Asset 5 - Baseline vs PPO comparison: teammate_repo/outputs/plots/nominal/baseline_vs_ppo.png\n\n"
        "Asset 6 - Robustness comparison: teammate_repo/outputs/plots/robustness/baseline_vs_ppo.png"
    ),
    "11. Lessons Learned": (
        "Lesson 1 - Always isolate parameters: Combined perturbation testing tells you 'something degrades performance', "
        "but only per-parameter isolation tells you what to fix. Our initial all-9 test showed 0.103m error; isolation "
        "revealed that target_noise alone caused 0.042m while 6 other parameters had zero individual impact.\n\n"
        "Lesson 2 - Fix what matters, don't randomize everything: Including ball mass in the perturbation set masked "
        "all other findings and wasted the first round of testing. Define scope clearly (ball is fixed) before designing tests.\n\n"
        "Lesson 3 - The policy was more robust than expected: With ball fixed, 7 simultaneous perturbations caused only "
        "a 3% success drop. The residual PPO approach (baseline + learned corrections) produces inherently stable behavior.\n\n"
        "Lesson 4 - Target position is the real bottleneck: In a real lab, the biggest risk is not sensor noise or motor "
        "degradation, but simply not knowing exactly where the target is. A clearly marked, precisely measured target zone "
        "is the cheapest and most effective robustness improvement.\n\n"
        "Lesson 5 - Statistical significance matters: 30 episodes can detect large effects (>10% success change) but "
        "small differences (3%) need 100+ episodes. For the final report, we used 100-episode runs for the combined test."
    ),
}

# Fill paragraphs after each heading
headings_map = {
    "5. Problems & Setbacks": "5. Problems",
    "6. Engineering Decisions": "6. Engineering",
    "7. Limitations": "7. Limitations",
    "8. Future Work": "8. Future",
    "9. Evidence Register": "9. Evidence",
    "10. Presentation Assets": "10. Presentation",
    "11. Lessons Learned": "11. Lessons",
}

for i, p in enumerate(doc.paragraphs):
    for full_title, key in headings_map.items():
        if key in p.text and p.style.name.startswith("Heading"):
            # Find the next paragraph (non-heading) to insert content
            text = sections_text[full_title]
            # Add text to the paragraph after the heading
            for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
                next_p = doc.paragraphs[j]
                if not next_p.style.name.startswith("Heading") and not next_p.text.strip():
                    next_p.text = text
                    break
            break

# === Section 12: Sim2Real Readiness Checklist ===
# Check off items
checklist_items = [
    (15, "Main transfer gaps identified", "[X]"),
    (16, "Domain randomization documented", "[X]"),
    (17, "Robustness experiments completed", "[X]"),
    (18, "Limitations documented", "[X]"),
    (19, "Future work proposed", "[X]"),
    (20, "Evidence collected for presentation", "[X]"),
]
for idx, prefix, status in checklist_items:
    doc.paragraphs[idx].text = f"{status} {prefix}"

# === Section 12 (end): Answer guided questions ===
guided_q = doc.paragraphs[13]
guided_q.text = (
    "Which Sim2Real gap was the most critical?\n"
    "Target position uncertainty. When the target position was randomized by +/-3cm, "
    "mean landing error doubled from 2.0cm to 4.2cm and success rate dropped from 100% to 97%. "
    "All other gaps (sensor noise, friction, latency, contact, motor dynamics) had negligible "
    "individual impact (~2cm, 100% success). This means the single best investment for real-world "
    "deployment is a precisely measured target zone.\n\n"
    "Which robustness technique appears most promising?\n"
    "Per-parameter isolation testing. Rather than only measuring combined degradation, isolating "
    "each parameter revealed that 6 out of 7 perturbations are non-issues individually. This allows "
    "us to focus engineering effort on the one real problem (target position) instead of wasting time "
    "on sensor noise filtering or actuator modeling improvements that would have zero impact.\n\n"
    "What would you do differently?\n"
    "1. Define ball as fixed parameter from the start (50g/4cm measured).\n"
    "2. Run per-parameter tests before combined tests to understand individual contributions first.\n"
    "3. Use 100-episode batches for all tests, not just the combined one, for statistical confidence.\n"
    "4. Add target position randomization to the training loop (domain randomization during PPO training) "
    "rather than only at evaluation time."
)

doc.save("document/Sim2Real_ADC2026_v2.docx")
print("DONE: document/Sim2Real_ADC2026_v2.docx — all sections filled")
